"""
Build the professional UC2_v2_Report.docx using python-docx.
Writes to UC2_v2/docs/UC2_v2_Report.docx.
"""

import os
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Inches, Pt, RGBColor


# Resolve project root portably. Order of preference:
#   1. UC2_ROOT environment variable (explicit override)
#   2. The script's own parent directory if it contains UC2_v2/
#   3. A few well-known candidate locations (sandbox mount + user Desktop)
def _resolve_project_root() -> Path:
    env = os.environ.get("UC2_ROOT")
    if env:
        p = Path(env).expanduser().resolve()
        if (p / "docs").exists() or (p / "UC2_v2").exists():
            return p if p.name == "UC2_v2" else p / "UC2_v2"
    here = Path(__file__).resolve().parent
    for candidate in (here, here.parent, here.parent.parent):
        if (candidate / "UC2_v2" / "docs").exists():
            return candidate / "UC2_v2"
        if candidate.name == "UC2_v2" and (candidate / "docs").exists():
            return candidate
    for candidate in (
        Path("/sessions/upbeat-stoic-ptolemy/mnt/Cap proj/UC2_v2"),
        Path.home() / "Desktop" / "Cap proj" / "UC2_v2",
    ):
        if candidate.exists():
            return candidate
    raise FileNotFoundError(
        "Could not locate UC2_v2/. Set UC2_ROOT to the UC2_v2 directory."
    )


ROOT = _resolve_project_root()
OUT = ROOT / "docs" / "UC2_v2_Report.docx"
OUT.parent.mkdir(parents=True, exist_ok=True)
FIG = OUT.parent / "figures"


def add_figure(path, width_inches=6.3, caption=None):
    """Insert a centred image with an optional italic caption below."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if path.exists():
        run = p.add_run()
        run.add_picture(str(path), width=Inches(width_inches))
    else:
        r = p.add_run(f"[missing figure: {path.name}]")
        r.italic = True
    if caption:
        c = doc.add_paragraph()
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = c.add_run(caption)
        r.italic = True
        r.font.size = Pt(9)

# -----------------------------------------------------------------------------
# Setup
# -----------------------------------------------------------------------------

doc = Document()

# Page: US Letter, 1-inch margins
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1.0)
section.bottom_margin = Inches(1.0)
section.left_margin = Inches(1.0)
section.right_margin = Inches(1.0)

# Base font
style_normal = doc.styles["Normal"]
style_normal.font.name = "Calibri"
style_normal.font.size = Pt(11)
style_normal.paragraph_format.space_after = Pt(6)
style_normal.paragraph_format.line_spacing = 1.15

for h_name, size in (("Heading 1", 18), ("Heading 2", 14), ("Heading 3", 12)):
    st = doc.styles[h_name]
    st.font.name = "Calibri"
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.color.rgb = RGBColor(0x1F, 0x3A, 0x68)
    st.paragraph_format.space_before = Pt(12 if h_name == "Heading 1" else 10)
    st.paragraph_format.space_after = Pt(6)


def _style_cell(cell, fill_hex, color="BFBFBF", sz="4"):
    """Set borders + shading in schema-required order (tcBorders before shd)."""
    tcPr = cell._tc.get_or_add_tcPr()

    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz)
        el.set(qn("w:color"), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)

    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def add_para(text, *, bold=False, italic=False, size=None, align=None, after=None, before=None, color=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    if before is not None:
        p.paragraph_format.space_before = Pt(before)
    if after is not None:
        p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    if size:
        r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    return p


def add_bullet(text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 + 0.25 * level)
    p.paragraph_format.space_after = Pt(3)
    p.add_run(text)
    return p


def add_table(headers, rows, col_widths=None, header_fill="1F3A68", zebra="F2F6FB"):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False

    if col_widths is None:
        total = Inches(6.5)
        col_widths = [Inches(6.5 / len(headers))] * len(headers)
    for i, w in enumerate(col_widths):
        for cell in table.columns[i].cells:
            cell.width = w

    # Header row
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        cell = hdr[i]
        cell.width = col_widths[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(h)
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.size = Pt(10.5)
        _style_cell(cell, header_fill)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for idx, row in enumerate(rows):
        rc = table.add_row().cells
        fill = zebra if idx % 2 == 0 else "FFFFFF"
        for i, val in enumerate(row):
            cell = rc[i]
            cell.width = col_widths[i]
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(str(val))
            r.font.size = Pt(10.5)
            _style_cell(cell, fill)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Spacing after table
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


# -----------------------------------------------------------------------------
# Title page
# -----------------------------------------------------------------------------

for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("UC2 — Inspector-Triggered\nTicket-Purchase Fraud Detection")
r.bold = True
r.font.size = Pt(28)
r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x68)

doc.add_paragraph()

add_para(
    "Hidden Markov Model pipeline for surfacing mobile-ticket accounts whose "
    "activation behaviour is consistent with inspector-triggered purchases.",
    italic=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER,
)

# Stamp that the document is built from the outputs directory
import json
SENS = FIG / "sensitivity_summary.json"
sens = None
if SENS.exists():
    try:
        sens = json.loads(SENS.read_text())
    except Exception:
        sens = None

for _ in range(10):
    doc.add_paragraph()

add_para("Final Project Report", bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para("WPI MS Fintech Capstone — Team 2 (Masabi × Gemsen)",
         size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para("April 2026", size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para(
    "Acknowledgements: Charlie Ko (Masabi) and LF Arsenault (Gemsen) for "
    "sponsor guidance, data access, and review comments that shaped the "
    "HMM-seed sweep, weight-sensitivity analysis, and symbol-vocabulary "
    "choices documented here.",
    italic=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER,
)

doc.add_page_break()

# -----------------------------------------------------------------------------
# Executive summary
# -----------------------------------------------------------------------------

doc.add_heading("Executive Summary", level=1)
add_para(
    "This project builds an end-to-end fraud-detection pipeline that identifies "
    "mobile-ticket accounts whose activation behaviour is consistent with "
    "inspector-triggered ticket purchases — accounts that buy a ticket only when a "
    "fare inspector is about to check it. The pipeline ingests approximately 20 GB of "
    "raw activation, purchase, and validation-scan logs (6.4 million activations across "
    "232,669 unique riders), derives a compact symbolic representation of each rider's "
    "activation stream, fits a Hidden Markov Model (HMM) to those symbol sequences, and "
    "produces two complementary human-review shortlists."
)
add_para(
    "The primary shortlist combines heuristic-rule violations with a model-based "
    "anomaly score. A supplementary shortlist isolates riders the heuristic rules would "
    "have missed entirely, so the incremental value of the HMM can be measured directly. "
    "On the full dataset, the HMM surfaces 88,328 rule-clean riders as candidates, and "
    "the top-100 among those spend essentially 100 % of their activation sequence in "
    "the model's high-risk state set (median posterior dominance 0.9999, maximum "
    "1.0000). This is direct, quantitative evidence that the model adds genuine "
    "discovery capability beyond what the heuristic rules catch."
)

# -----------------------------------------------------------------------------
# 1 Problem statement
# -----------------------------------------------------------------------------

doc.add_heading("1. Problem Statement", level=1)
add_para(
    "In a mobile-ticketing system, a rider is expected to activate their ticket before "
    "boarding. Fare inspectors typically scan the activated ticket using a handheld "
    "device within seconds of boarding. A common fraud pattern is the inspector-"
    "triggered purchase: the rider buys and activates a ticket only when they see an "
    "inspector approaching, producing a very short interval between purchase, "
    "activation, and scan. At scale, this pattern is indistinguishable from a "
    "legitimate late activation on any single event, but is highly distinctive across a "
    "rider's full activation history."
)
add_para(
    "The project's goal is to surface the small subset of accounts whose activation "
    "history is most consistent with this behaviour pattern, in a form suitable for "
    "human review. Two complementary approaches are combined: a heuristic rule layer "
    "that captures the narrowest, most obvious cases, and a probabilistic sequence "
    "model that learns the latent state structure behind each rider's behaviour and "
    "scores them by how often they occupy high-risk states."
)

# -----------------------------------------------------------------------------
# 2 Dataset
# -----------------------------------------------------------------------------

doc.add_heading("2. Dataset", level=1)
add_para(
    "The pipeline consumes four primary operational tables and three enrichment "
    "reference tables. All primary tables are timestamped in UTC; the loader "
    "layer validates tz-awareness on every timestamp series before any downstream "
    "feature is computed."
)
add_para(
    "Timestamp source. All gap calculations (the HIGH 15-second window, the "
    "MEDIUM 30-second window, the activation-to-handheld intervals that emit "
    "the HMM symbols, and the purchase-to-activation intervals) use the "
    "server_timestamp field from each raw CSV — the timestamp the Masabi "
    "backend assigned when the event was received. Server timestamps are the "
    "authoritative clock for gap arithmetic: they come from a single "
    "operational source, whereas device-local timestamps (including the "
    "activation_timestamp column denormalised onto the validation_scans "
    "table) can drift from the server clock or be manipulated by a "
    "compromised device, and that drift would shift infraction counts "
    "substantially at the 15-second boundary. The denormalised "
    "activation_timestamp on validation_scans is retained only to "
    "cross-reference which ticket a scan validates; it never enters the "
    "HIGH / MEDIUM window arithmetic."
)

add_table(
    headers=["Table", "Size", "Rows", "Role"],
    rows=[
        ["retail_activations.csv", "5.2 GB", "6,407,862", "primary activation log"],
        ["retail_ticket_purchases.csv", "2.9 GB", "3,530,890", "primary purchase log"],
        ["retail_tickets.csv", "6.0 GB", "6,464,116", "bridge table (reference)"],
        ["validation_scans.csv", "6.1 GB", "6,371,275", "handheld / gate validations"],
        ["calendar_of_events.csv", "—", "—", "service disruption & event flags"],
        ["commuter_rail_stops.csv", "—", "—", "station reference"],
        ["commuter_rail_boardings_by_line.csv", "—", "—", "line-level volume"],
    ],
    col_widths=[Inches(2.2), Inches(0.8), Inches(1.3), Inches(2.2)],
)

add_para(
    "After dropping rows with null account identifiers, 3,679,391 activation events "
    "remain across 232,669 unique riders. 98,241 riders clear the minimum-five-"
    "activation floor required for HMM training. Calendar enrichment contributes five "
    "per-rider exposure features (mean exposure to schedule-impact, maintenance-impact, "
    "event-impact, school-impact, and holiday-impact days) so that legitimate gap-"
    "looking behaviour on disruption days is not misread as anomalous."
)
add_para(
    "Deviations from the Data Readiness checklist. Three items from the sponsor "
    "data-readiness report are worth surfacing explicitly so the downstream numbers "
    "are read in the right context. First, the account_id null rate on "
    "retail_activations is approximately 43 % of raw rows (6,407,862 raw → "
    "3,679,391 kept), well above the 5 % threshold flagged as a soft concern in "
    "the readiness checklist. Operationally, this is consistent with legitimate "
    "guest / anonymous activation paths rather than a data-export defect, and the "
    "remaining 3.68 M identified rows still exceed the volume required for HMM "
    "training; the pipeline documents the drop rate in RUN_RESULTS.md so a future "
    "reviewer can challenge it. Second, station exclusion via the scanned_at JSON "
    "field (Data Readiness Activity 6, non-blocker) is not implemented in this "
    "version: the validation_scans table in the current export ships 100 % "
    "handheld records, so there are no gate-station scans to exclude and the "
    "rule-window counts are numerically unchanged from what a station-aware "
    "implementation would return. Third, the commuter_rail_stops and "
    "commuter_rail_boardings_by_line enrichment tables are loaded by uc2_io but "
    "are not joined at the per-rider level, because the activation events in "
    "this export carry no station_id or line_id — these two tables are retained "
    "as reference inputs rather than features for this run and are candidates "
    "for a follow-up join if a geo-tagged activation feed becomes available."
)

# -----------------------------------------------------------------------------
# 3 Methodology
# -----------------------------------------------------------------------------

doc.add_heading("3. Methodology", level=1)

# --- 3.1 Symbol vocabulary
doc.add_heading("3.1 Activation Symbol Vocabulary", level=2)
add_para(
    "Each activation event is mapped to exactly one of seven symbols based on what "
    "happens in the surrounding time window. The mapping is first-match-wins: the "
    "emission engine evaluates rules in order, and a rider who purchases and then "
    "activates within 60 seconds followed by a fast handheld scan is tagged "
    "PURCHASE_THEN_ACTIVATE_FAST, not ACTIVATE_FAST_HANDHELD. This ordering preserves "
    "the inspector-triggered signal the pipeline is trying to isolate."
)

add_table(
    headers=["ID", "Symbol", "Emission rule"],
    rows=[
        ["0", "ACTIVATE_FAST_HANDHELD",      "handheld scan within 15 s of activation"],
        ["1", "ACTIVATE_GAMING_THRESHOLD",   "handheld scan 16 – 30 s after activation"],
        ["2", "ACTIVATE_SLOW_HANDHELD",      "handheld scan 30 – 300 s after activation"],
        ["3", "ACTIVATE_GATE",               "gate scan within 120 s, no fast handheld"],
        ["4", "NO_HANDHELD_FOLLOWUP",        "no handheld or gate scan in window"],
        ["5", "OTHER_HANDHELD_PATTERN",      "fall-through catch-all"],
        ["6", "PURCHASE_THEN_ACTIVATE_FAST", "purchase ≤ 60 s before, handheld ≤ 15 s after"],
    ],
    col_widths=[Inches(0.5), Inches(2.4), Inches(3.6)],
)

add_para(
    "The 15-second boundary for ACTIVATE_FAST_HANDHELD is chosen deliberately to match "
    "the HIGH rule window (Section 3.2), which flags clusters of gaps of 15 seconds or "
    "less. Aligning the symbol cutoff with the heuristic threshold means the rule layer "
    "and the HMM layer fire on the same physical event, so rule-based evidence and "
    "model-based evidence corroborate each other rather than measuring slightly "
    "different phenomena. The 16 – 30 second GAMING_THRESHOLD band then captures the "
    "activation behaviour that stays just outside the rule window — the dodge pattern "
    "that a rider would adopt if they were trying to evade a strict 15-second rule — "
    "and lets the HMM weight that behaviour through its emission matrix without "
    "hard-coding a penalty."
)
add_para(
    "Range of ACTIVATE_SLOW_HANDHELD. The 30 – 300 second SLOW band subsumes both "
    "the normal inspector-encounter window (roughly 30 – 120 s, where a rider is "
    "flagged down and produces a ticket at normal pace) and the very-slow tail "
    "(120 – 300 s, where the scan is effectively unrelated to the activation). "
    "A split into SLOW (30 – 120 s) and VERY_SLOW (> 120 s) was evaluated as a "
    "reviewer-suggested alternative, but the unified band already distinguishes "
    "the fraud-shaped signal (FAST, GAMING, PURCHASE_THEN_FAST) from routine "
    "handheld activity and preserves interpretability at 7 symbols. If a future "
    "run wants to separate these two sub-bands, the split is isolated to "
    "src/uc2_symbols.py and does not touch the HMM training, scoring, or "
    "validation stages."
)

# --- 3.2 Pattern rules
doc.add_heading("3.2 Rule-Based Pattern Windows", level=2)
add_para(
    "Two heuristic rules label riders whose activation stream contains clustered fast "
    "gaps. A rider satisfies the HIGH pattern when three or more activation gaps of 15 "
    "seconds or less occur within any sliding 240-hour window, and the MEDIUM pattern "
    "when three or more gaps of 30 seconds or less occur within any sliding 168-hour "
    "window. Both counts are computed directly from the sorted per-account activation "
    "series using an O(n) two-pointer sweep, so the counts are exact by construction "
    "and repeat_offender_flag is a pure function of max_infractions_240h that cannot "
    "drift from the window definition."
)

# --- 3.3 HMM training
doc.add_heading("3.3 Hidden Markov Model Training", level=2)
add_para(
    "A Categorical HMM is fit over the rider-level symbol sequences. The training "
    "sweep spans three state counts (7, 9, 11) and eight random seeds (0 – 7), yielding "
    "24 independent fits. Fits run in parallel through a ProcessPoolExecutor using the "
    "fork multiprocessing context on two-thirds of the available CPU cores. Model "
    "selection uses the Bayesian Information Criterion (BIC), which balances "
    "log-likelihood against parameter count and observation count — the fit with the "
    "minimum BIC is selected and its emission matrix is written out for downstream "
    "labelling of high-risk states. When hmmlearn is unavailable, a pure-numpy "
    "Baum-Welch implementation serves as a drop-in fallback."
)
add_para(
    "Why eight seeds. Baum-Welch is an expectation-maximisation algorithm, which "
    "is gradient-style and has no guarantee of finding the global optimum: a "
    "7-state model with an unfortunate initialisation can fit worse than an "
    "11-state model with a lucky one, purely because EM got stuck in a local "
    "minimum on the first run. Single-seed sweeps therefore measure "
    "initialisation luck rather than model capacity. Running eight independent "
    "seeds per state count and selecting the minimum BIC across all 24 fits "
    "defends against this failure mode — the selected fit sits 31,000 BIC "
    "points below the nearest runner-up, a gap that would not be visible from "
    "a single-seed run. The hmmlearn convergence flag (which by design returns "
    "True whenever the fit exhausts its iteration budget) is not used as a "
    "stopping criterion for model selection; fifty iterations per seed is "
    "sufficient to reach the likelihood plateau, past which seed diversity "
    "matters more than iteration depth."
)

# --- 3.4 Eligibility
doc.add_heading("3.4 Eligibility and Sequence Preparation", level=2)
add_para(
    "Only riders with at least five activation events enter HMM training. Each rider's "
    "symbol stream is capped FIFO at the 30 most-recent symbols, so that long-tenured "
    "accounts with hundreds of routine activations do not swamp the likelihood and "
    "wash out the gaming-threshold and gate-scan states."
)

# --- 3.5 Anomaly scoring
doc.add_heading("3.5 Anomaly Scoring", level=2)
add_para(
    "Riders are ranked primarily by posterior dominance. For each rider, the forward-"
    "backward algorithm produces smoothed state posteriors across their sequence. The "
    "posterior-dominance score is the mean mass those posteriors place on a post-hoc-"
    "selected high-risk state set. High-risk states are identified by inspecting the "
    "fitted emission matrix: states whose emission distribution concentrates on fraud-"
    "shaped symbols (ACTIVATE_FAST_HANDHELD, ACTIVATE_GAMING_THRESHOLD, "
    "PURCHASE_THEN_ACTIVATE_FAST) are labelled high-risk, and the top half of states "
    "by that criterion form the high-risk set."
)
add_para(
    "The final combined anomaly score blends four normalised components:"
)
add_bullet("Posterior dominance — weight 0.50")
add_bullet("Rule-violation count — weight 0.30")
add_bullet("Gaming-band ratio (fraction of gaps in the 16 – 30 s dodge band) — weight 0.15")
add_bullet("Raw burst count — weight 0.05")
add_para(
    "A proportional penalty is then applied to burst-only riders — accounts with 50 "
    "or more bursts but fewer than 3 real infractions — so that short, repetitive "
    "activation bursts do not dominate the shortlist at the expense of hybrid "
    "burst-and-infraction behaviour."
)

# --- 3.6 Validation
doc.add_heading("3.6 Rule-vs-HMM Validation", level=2)
add_para(
    "Two independent shortlists quantify how the model and the rules interact. "
    "The primary shortlist is the top 100 riders by the combined anomaly score. The "
    "supplementary shortlist is the top 100 riders by posterior dominance among those "
    "with zero rule infractions. Notebook 04 reports the overlap (rule-confirmed "
    "riders, rule-only riders, HMM-only riders, and the supplementary pool size) so "
    "that the contributions of the two layers can be audited separately."
)

# -----------------------------------------------------------------------------
# 4 Implementation
# -----------------------------------------------------------------------------

doc.add_heading("4. Implementation", level=1)
add_para(
    "The pipeline is organised as a thin set of Jupyter notebooks over an importable, "
    "unit-testable Python layer. Every non-trivial operation lives in a module under "
    "src/ so the notebooks can be read as a narrative while the logic is reviewed in "
    "isolation."
)

add_table(
    headers=["Module", "Responsibility"],
    rows=[
        ["src/uc2_io.py",         "CSV readers, chunked loading, UTC validation, enrichment joins"],
        ["src/uc2_symbols.py",    "seven-symbol vocabulary and the emit_symbol engine"],
        ["src/uc2_features.py",   "pattern windows, sequence preparation, timing aggregates"],
        ["src/uc2_hmm_utils.py",  "parallel multi-seed / multi-state HMM training and selection"],
        ["src/uc2_scoring.py",    "posterior-dominance scoring and burst de-weight"],
    ],
    col_widths=[Inches(1.8), Inches(4.7)],
)

add_para(
    "The four notebooks run in order and pass artefacts through parquet, pickle, and "
    "npz files in the outputs/ directory:"
)
add_bullet("01_UC2_Feature_Engineering.ipynb — raw CSVs to feature table, symbol rows, and HMM sequences")
add_bullet("02_UC2_HMM_Training.ipynb — symbol sequences to fitted HMM plus BIC sweep")
add_bullet("03_UC2_Exercise3_Scoring.ipynb — fitted HMM to rider scores and primary shortlist")
add_bullet("04_UC2_Rule_Based_Validation.ipynb — rule-vs-HMM overlap and supplementary shortlist")

add_para(
    "The 20 GB working set is handled on a 16 GB laptop by reading each CSV in 1 "
    "million-row chunks, casting identifier columns to categorical dtype, and casting "
    "only back to object at the merge boundary (pandas merge_asof refuses to merge "
    "two categorical columns with differing category sets). Peak resident memory on "
    "the full dataset is approximately 7 GB, leaving 9 GB headroom."
)

# -----------------------------------------------------------------------------
# 5 Results
# -----------------------------------------------------------------------------

doc.add_heading("5. Results", level=1)
add_para(
    "All numbers below come from a single end-to-end run of the pipeline on the full "
    "dataset. Total wall-clock time was approximately 117 minutes on a 16 GB M-series "
    "laptop (22 min feature engineering, 90 min HMM training grid at 8 seeds × 3 "
    "state counts, 3 min scoring, 2 min validation)."
)

doc.add_heading("5.1 Feature Engineering", level=2)
add_table(
    headers=["Quantity", "Value"],
    rows=[
        ["Raw activation rows",                                     "6,407,862"],
        ["Activation events (after dropna on account_id)",          "3,679,391"],
        ["Validation scans",                                        "6,371,275 handheld / 0 gate"],
        ["Purchase events kept",                                    "≈ 2.0 M (from 3,530,890 raw)"],
        ["Unique riders",                                           "232,669"],
        ["HMM-eligible riders (≥ 5 activations)",                   "98,241"],
        ["Rule-flagged riders (HIGH or MEDIUM)",                    "10,708"],
        ["  HIGH only (240 h window, ≥ 3 gaps ≤ 15 s)",             "9,115"],
        ["  MEDIUM only (168 h window, ≥ 3 gaps ≤ 30 s)",           "10,427"],
    ],
    col_widths=[Inches(4.2), Inches(2.3)],
)

doc.add_heading("5.2 HMM Training", level=2)
add_table(
    headers=["Metric", "Value"],
    rows=[
        ["Training observations",                        "1,767,777"],
        ["Training sequences (= HMM-eligible riders)",   "98,241"],
        ["Sequence length p50 / p90 / p99",              "16 / 30 / 30"],
        ["Selected state count (minimum BIC)",           "9"],
        ["Best seed",                                    "6"],
        ["Log-likelihood",                               "−2,640,637.1"],
        ["BIC (selected)",                               "5,283,200.8"],
        ["Runner-up (11 states, seed 5) BIC",            "5,314,340.0"],
    ],
    col_widths=[Inches(3.6), Inches(2.9)],
)

add_figure(
    FIG / "bic_sweep.png",
    width_inches=6.3,
    caption="Figure 1. BIC across all 24 fits (8 seeds × 3 state counts). "
            "The selected fit (red) sits 31,000 points below the nearest runner-up, "
            "a margin that would have been missed by the original 5-seed sweep.",
)

add_figure(
    FIG / "emissions_heatmap.png",
    width_inches=6.3,
    caption="Figure 2. Fitted emission matrix (9 states × 7 symbols). "
            "Red-outlined rows are the four high-risk states identified post-hoc "
            "from emission mass on fraud-shaped symbols "
            "(FAST, GAMING, PURCH+FAST). State 0 and State 7 are the strongest "
            "fast-scanner states; State 3 is the purchase-then-fast state. "
            "State 6 (99 % NO_FU) captures silent riders with no scan follow-up.",
)

doc.add_heading("5.3 Primary Shortlist (Combined Score)", level=2)
add_table(
    headers=["Metric", "Value"],
    rows=[
        ["Riders scored",                                   "98,241"],
        ["Burst-only riders in top-100",                    "0"],
        ["Rule-confirmed riders in top-100 (R ∩ H)",        "100 / 100"],
        ["HMM-only riders in top-100 (H ∖ R)",              "0"],
    ],
    col_widths=[Inches(3.6), Inches(2.9)],
)
add_para(
    "With 10,708 rule-flagged riders competing for 100 slots on a score that "
    "incorporates rule infractions, the combined score naturally promotes rule-flagged "
    "riders to the top. The 100-out-of-100 overlap is a methodology-agreement result: "
    "the HMM confirms the rule-based findings rather than manufacturing new false "
    "positives. The supplementary shortlist is where the incremental discovery is "
    "measured."
)

add_figure(
    FIG / "score_histogram.png",
    width_inches=6.3,
    caption="Figure 3. Combined anomaly score distribution across all 98,241 "
            "scored riders (log-scaled y-axis). The red dashed line marks the "
            "top-100 cutoff used for the primary shortlist; the shortlist sits "
            "at a distributional break rather than an arbitrary cut.",
)

doc.add_heading("5.4 Supplementary Shortlist (HMM-only discovery)", level=2)
add_table(
    headers=["Metric", "Value"],
    rows=[
        ["Non-rule-flagged scored pool",              "88,328"],
        ["Supplementary shortlist size",              "100"],
        ["Top posterior dominance",                   "1.0000"],
        ["Median posterior dominance",                "0.9999"],
    ],
    col_widths=[Inches(3.6), Inches(2.9)],
)
add_para(
    "88,328 riders had zero rule infractions yet still cleared the five-activation "
    "HMM eligibility floor. Among those, the top 100 selected by posterior dominance "
    "all spend essentially 100 % of their activations in the model's high-risk state "
    "set (median 0.9999, maximum 1.0000). These are the riders the heuristic rules "
    "would have missed entirely, and the model pins them at the ceiling — direct, "
    "quantitative evidence that the HMM surfaces a distinct and defensible population "
    "for human review."
)

doc.add_heading("5.5 Shortlist Stability Across Near-Winning Models", level=2)
add_para(
    "To test whether the top-100 primary shortlist depends on which specific fit "
    "the BIC selection lands on, the winning model was compared against the two "
    "lowest-BIC runner-ups (11 states / seed 5, BIC 5,314,340; and 11 states / "
    "seed 6, BIC 5,335,142). Each alternative model was refit from the same "
    "symbol sequences, then a top-100 shortlist was produced using identical "
    "scoring weights. Pairwise Jaccard overlap between the three shortlists "
    "ranged from 0.36 to 0.44."
)
add_table(
    headers=["Pair", "Jaccard overlap", "Shared top-100 riders"],
    rows=[
        ["Winner (#1) vs. runner-up #2 (11 states, seed 5)",  "0.44",  "≈ 61 / 100"],
        ["Winner (#1) vs. runner-up #3 (11 states, seed 6)",  "0.36",  "≈ 53 / 100"],
        ["Runner-up #2 vs. runner-up #3",                     "0.43",  "≈ 60 / 100"],
    ],
    col_widths=[Inches(3.6), Inches(1.3), Inches(1.6)],
)
add_figure(
    FIG / "shortlist_stability.png",
    width_inches=5.6,
    caption="Figure 4. Pairwise Jaccard overlap between the top-100 shortlists "
            "produced by the winning model and the two lowest-BIC runner-ups.",
)
add_para(
    "Interpretation. The combined anomaly score weights rule-infraction count "
    "at 0.30, so every near-winning model surfaces top-100 shortlists that are "
    "drawn from the same 10,708-rider rule-flagged population. The observed "
    "Jaccard of 0.36 – 0.44 reflects model-specific reordering within that "
    "population rather than disagreement about which riders to investigate. "
    "In practical terms, a reviewer working through any of the three top-100 "
    "shortlists would see qualitatively similar accounts drawn from the same "
    "candidate pool. Because BIC selects the winning fit with a 31,000-point "
    "margin over the nearest runner-up, the winning model's shortlist is the "
    "authoritative ranking; the stability check simply establishes that this "
    "ranking is not a knife-edge artefact of a single fit."
)

doc.add_heading("5.6 Weight Sensitivity of the Combined Score", level=2)
add_para(
    "The combined anomaly score is a weighted sum of four normalised "
    "components (posterior dominance 0.50, rule-violation count 0.30, "
    "gaming-band ratio 0.15, burst count 0.05). To confirm the top-100 "
    "shortlist is not an artefact of those specific weights, the combined "
    "score was recomputed under five alternative weight schemes — HMM-heavy "
    "(0.70 / 0.15 / 0.10 / 0.05), rule-heavy (0.20 / 0.60 / 0.15 / 0.05), "
    "equal weights (0.25 each), HMM-only (1.0 / 0 / 0 / 0), and rules-only "
    "(0 / 1.0 / 0 / 0) — and each scheme's top-100 shortlist was compared "
    "against the baseline shortlist by Jaccard overlap. The burst de-weight "
    "(0.25 multiplier on burst-only riders) is retained in every scheme so "
    "that the comparison isolates the effect of weight choice from the "
    "effect of the burst correction."
)

if sens:
    rows = []
    for row in sens["schemes"]:
        rows.append([
            row["scheme"],
            f"{row['jaccard_vs_baseline']:.2f}",
            f"{row['shared_with_baseline']} / 100",
        ])
    add_table(
        headers=["Scheme", "Jaccard vs. baseline", "Shared top-100 riders"],
        rows=rows,
        col_widths=[Inches(3.6), Inches(1.4), Inches(1.5)],
    )
    add_para(
        f"Median pairwise Jaccard across all six schemes: "
        f"{sens['median_pairwise_jaccard']:.2f} "
        f"(range {sens['min_pairwise_jaccard']:.2f} – "
        f"{sens['max_pairwise_jaccard']:.2f})."
    )
else:
    add_para(
        "Run make_sensitivity.py to populate this table. The script "
        "reads outputs/rider_scores.parquet, recomputes the top-100 "
        "shortlist under each scheme, writes docs/figures/"
        "sensitivity_summary.json, and the next rebuild of this "
        "document embeds the numbers here automatically.",
        italic=True,
    )

add_figure(
    FIG / "sensitivity.png",
    width_inches=6.0,
    caption="Figure 5. Pairwise Jaccard overlap between top-100 shortlists "
            "produced under six weighting schemes. The baseline weights "
            "(S1: 0.50 / 0.30 / 0.15 / 0.05) are the published choice; "
            "S2–S6 sweep the space from HMM-dominant to rule-dominant to "
            "equal-weight to single-component extremes.",
)

add_para(
    "Interpretation. The three blended schemes (HMM-heavy, rule-heavy, "
    "and equal-weights) share 53 to 74 of their top-100 riders with the "
    "baseline — substantial agreement given that each varies a single "
    "weight by a factor of two or more. The HMM-heavy scheme holds up "
    "best (74 / 100 shared) because the high-risk states the HMM learns "
    "are precisely the states rule-flagged riders spend most of their "
    "time in, so the two signals reinforce rather than contradict each "
    "other. The two single-component schemes are outliers by "
    "construction and are included as stress tests. Rules-only (9 / 100) "
    "collapses because rule-violation count cannot distinguish among the "
    "10,708 rule-flagged riders — most are tied at three to six "
    "infractions, so its top-100 is essentially arbitrary within that "
    "tied population. HMM-only (1 / 100) ignores rules entirely and "
    "surfaces the supplementary-shortlist population of rule-clean, "
    "posterior-dominant riders — by design the same population analysed "
    "in Section 5.4. The operational takeaway is that the baseline "
    "weights reflect an explicit policy choice — rules contribute 0.30 "
    "so that a rule-flagged rider with weak model signal is still "
    "surfaced — and reasonable perturbations of that choice produce "
    "top-100 shortlists that remain majority-consistent with the "
    "baseline."
)

doc.add_heading("5.7 Behavioural Findings", level=2)
add_para(
    "Two behavioural patterns surfaced while cross-validating the shortlists "
    "against the broader activation log and are worth recording alongside the "
    "model results, because they corroborate the latent-state structure the "
    "HMM learned and point at the real-world contexts in which "
    "inspector-triggered purchasing concentrates."
)
add_para(
    "Temporal concentration. Burst-shaped activation events are strongly "
    "non-uniform across the week. Wednesday evenings alone generate on the "
    "order of 425,000 short-interval activation events in the full export — "
    "roughly five times the Sunday baseline — matching the inspector-shift "
    "pattern reported by the sponsor and consistent with the HMM's decision "
    "to concentrate fast-scanner emission mass in States 0 and 7 rather than "
    "spread it uniformly. Riders surfaced on both shortlists are "
    "disproportionately active during these high-enforcement windows."
)
add_para(
    "Geographic concentration. Among riders who produce any ACTIVATE_FAST_"
    "HANDHELD symbol, a disproportionate share of activity concentrates on "
    "the Providence / Stoughton corridor: roughly 37 % of fast-event riders "
    "are associated with that line, and top-rider fast-event volume on it "
    "runs near five times the next-ranked line. Because the current "
    "activations feed does not carry a line_id column, this finding is drawn "
    "from cross-referencing the sponsor-provided behavioural presentation "
    "(Masabi x WPI x Gemsen weekly, 15 April 2026) rather than from a "
    "per-rider join inside this pipeline; surfacing it here flags the "
    "feature for the joined-geo follow-up referenced in Section 2."
)
add_para(
    "Latent-state corroboration. The emission matrix (Figure 2) shows State "
    "3 and State 6 as a bimodal pair — State 3 places 43 % of its emission "
    "mass on ACTIVATE_FAST_HANDHELD, while State 6 places 99 % on "
    "NO_HANDHELD_FOLLOWUP. The per-rider posterior mass on this "
    "suspicious-state cluster is itself bimodal across the eligible "
    "population: most riders sit near 0 %, a clear tail sits near 100 %, and "
    "the middle is sparsely populated. That shape is exactly what a fraud-"
    "detection shortlist needs — a well-separated high-risk mode, rather "
    "than a continuous smear — and it is the reason posterior dominance can "
    "be used as a ranking primitive without an arbitrary probability cutoff."
)

# -----------------------------------------------------------------------------
# 6 Deliverables
# -----------------------------------------------------------------------------

doc.add_heading("6. Deliverables", level=1)
add_para("All artefacts land in UC2_v2/outputs/ and are consumed by the two CSV shortlists at the bottom of the list.")
add_bullet("feature_table.parquet / .pkl — rider × feature matrix with calendar flags")
add_bullet("symbol_rows.parquet / .pkl — long-format symbol rows")
add_bullet("sequences.npz — HMM training input (≥ 5-event riders, FIFO-capped at 30)")
add_bullet("hmm_best.pkl — fitted CategoricalHMM plus BIC grid")
add_bullet("hmm_emissions.csv — K × V emission matrix")
add_bullet("hmm_grid_results.csv — full BIC / AIC sweep across all fits")
add_bullet("rider_scores.parquet / .pkl — per-rider combined score and components")
add_bullet("uc2_human_review_shortlist_v2.csv — primary top-100 combined shortlist")
add_bullet("uc2_hmm_only_riders.csv — supplementary top-100 HMM-only shortlist")
add_bullet("uc2_rule_vs_hmm_overlap.csv — R / H / R ∩ H / H ∖ R / supplementary pool counts")

# -----------------------------------------------------------------------------
# 7 Reproducibility
# -----------------------------------------------------------------------------

doc.add_heading("7. Reproducibility", level=1)
add_para(
    "Every notebook is deterministic given a fixed random seed. The HMM training grid "
    "uses seeds 0 – 7 explicitly; the parallel executor introduces non-determinism "
    "only into log-line ordering, not into model selection. On any machine with "
    "pandas 2.0+, numpy 1.24+, pyarrow 11+, and hmmlearn 0.3+ installed, the full "
    "pipeline reproduces the numbers in Section 5 exactly."
)
add_para(
    "The README describes how the data directory is discovered automatically at "
    "runtime, and RUN_RESULTS.md records per-table memory budgets and wall-time "
    "measurements from the latest run, including fallback strategies for machines "
    "with less than 16 GB of RAM."
)

# -----------------------------------------------------------------------------
# 8 Conclusion
# -----------------------------------------------------------------------------

doc.add_heading("8. Conclusion", level=1)
add_para(
    "The pipeline closes the loop from raw operational logs to a ranked, auditable "
    "human-review list in a fully reproducible form. The heuristic rules catch the "
    "narrow, obvious cases; the HMM extends reach by scoring every eligible rider on "
    "how often their activation sequence occupies high-risk latent states. On the full "
    "dataset, the model agrees with the rules where they overlap (100 / 100 on the "
    "combined top-100) and, critically, identifies an additional 100 riders with "
    "near-unit posterior dominance whom the rules would not have flagged. The two "
    "shortlists together give reviewers a complete picture: confirmed high-signal "
    "accounts plus the new candidates the rules would have let through."
)
add_para(
    "Several parameters remain available for operational tuning: the weighting of the "
    "four anomaly-score components, the 0.3 posterior-dominance cutoff for the "
    "posterior-driven label on the primary shortlist, the burst de-weight factor of "
    "0.25, the 120-second window for gate-scan emission, and the choice to take the "
    "top half of states as high-risk. All are isolated in src/uc2_scoring.py and "
    "src/uc2_symbols.py and can be adjusted without touching the notebooks."
)

# -----------------------------------------------------------------------------
# Save
# -----------------------------------------------------------------------------

doc.save(str(OUT))

# -----------------------------------------------------------------------------
# Fix validation: ensure w:zoom has a percent attribute
# -----------------------------------------------------------------------------
import zipfile, shutil, re

tmp = OUT.with_suffix(".tmp.docx")
with zipfile.ZipFile(OUT, "r") as zin, zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
    for item in zin.namelist():
        data = zin.read(item)
        if item == "word/settings.xml":
            text = data.decode("utf-8")
            text = re.sub(
                r'<w:zoom(\s+[^/>]*)?\s*/>',
                lambda m: '<w:zoom w:percent="100"/>',
                text,
            )
            # handle self-closing without attributes or with attrs but no percent
            if 'w:percent' not in text:
                text = text.replace('<w:zoom/>', '<w:zoom w:percent="100"/>')
            data = text.encode("utf-8")
        zout.writestr(item, data)
shutil.move(str(tmp), str(OUT))

print(f"Wrote {OUT} ({OUT.stat().st_size:,} bytes)")
